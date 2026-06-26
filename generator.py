"""
Generates a professional PRCE Word document from PRCEMetrics.
Styled to match A.Ramalhão's template (teal headers, structured tables).
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime
from calculator import PRCEMetrics

# ── Brand colors (hex strings for XML + RGBColor for run.font.color) ─────────
TEAL_HEX     = "1F778F"
DARK_BLUE_HEX = "1A3A5C"
WHITE_HEX    = "FFFFFF"
LIGHT_BG_HEX = "F2F7FA"
WARN_BG_HEX  = "FFF3CD"   # amber warning background

TEAL      = RGBColor(0x1F, 0x77, 0x8F)
DARK_BLUE = RGBColor(0x1A, 0x3A, 0x5C)
WHITE     = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT_BG  = RGBColor(0xF2, 0xF7, 0xFA)


def _set_cell_bg(cell, hex_color: str):
    """Set cell background using a hex string like 'F2F7FA'."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _header_row(table, texts: list[str], bg_hex=TEAL_HEX, fg=WHITE, bold=True):
    row = table.rows[0]
    for i, text in enumerate(texts):
        if i < len(row.cells):
            cell = row.cells[i]
            _set_cell_bg(cell, bg_hex)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(text)
            run.bold = bold
            run.font.color.rgb = fg
            run.font.size = Pt(9)


def _data_row(table, row_idx: int, values: list, bold_first=False, bg_hex=None):
    row = table.rows[row_idx]
    for i, val in enumerate(values):
        if i < len(row.cells):
            cell = row.cells[i]
            if bg_hex:
                _set_cell_bg(cell, bg_hex)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(str(val))
            run.font.size = Pt(9)
            if bold_first and i == 0:
                run.bold = True


def _add_section_heading(doc, number: str, title: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    run_num = p.add_run(f"{number}. ")
    run_num.bold = True
    run_num.font.color.rgb = TEAL
    run_num.font.size = Pt(11)
    run_title = p.add_run(title.upper())
    run_title.bold = True
    run_title.font.color.rgb = TEAL
    run_title.font.size = Pt(11)


def _add_table_caption(doc, number: str, title: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    run = p.add_run(f"Tabela {number} - {title}")
    run.bold = True
    run.font.size = Pt(9)


def _fmt_num(val, decimals=1):
    if val is None:
        return "N/A"
    if isinstance(val, (int, float)):
        return f"{val:,.{decimals}f}".replace(",", "X").replace(".", ",").replace("X", ".")
    return str(val)


def generate_prce(metrics: PRCEMetrics, company_name: str = "XXXX", output_path: str = None) -> str:
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # Default paragraph font
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(10)

    import re as _re
    year_match = _re.search(r"\b(20\d{2})\b", metrics.data_period)
    ref_year = year_match.group(1) if year_match else "2024"
    plan_start = int(ref_year) + 1
    plan_end   = plan_start + 2
    today = datetime.date.today().strftime("%d/%m/%Y")

    # ── COVER PAGE ───────────────────────────────────────────────────────────
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Cm(3)
    run = p.add_run("PLANO DE RACIONALIZAÇÃO DO\nCONSUMO DE ENERGIA")
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = DARK_BLUE

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(f"{plan_start} - {plan_end}")
    r2.font.size = Pt(16)
    r2.font.color.rgb = TEAL

    doc.add_paragraph()

    # Info box
    table_cover = doc.add_table(rows=4, cols=1)
    table_cover.style = "Table Grid"
    table_cover.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_cell_bg(table_cover.rows[0].cells[0], LIGHT_BG_HEX)
    _set_cell_bg(table_cover.rows[1].cells[0], LIGHT_BG_HEX)
    _set_cell_bg(table_cover.rows[2].cells[0], LIGHT_BG_HEX)
    _set_cell_bg(table_cover.rows[3].cells[0], LIGHT_BG_HEX)

    def cover_line(row_idx, label, value="", label_color=TEAL):
        cell = table_cover.rows[row_idx].cells[0]
        p = cell.paragraphs[0]
        r = p.add_run(label)
        r.bold = True
        r.font.color.rgb = label_color
        r.font.size = Pt(10)
        if value:
            r2 = p.add_run(f"  {value}")
            r2.font.size = Pt(10)

    cover_line(0, company_name)
    cover_line(1, "Frota de Transportes")
    cover_line(2, f"Ano de Referência: {ref_year}")
    cover_line(3, f"PRCE_{ref_year}_V0")

    doc.add_paragraph()

    meta = doc.add_paragraph()
    for label in ["Elaborado por:", "Equipa Técnica:", "Projeto:", "Versão: 0", f"Data: {today}"]:
        r = meta.add_run(label + "\n")
        r.bold = True
        r.font.size = Pt(9)

    doc.add_page_break()

    # ── PARTIAL DATA WARNING ─────────────────────────────────────────────────
    notes = getattr(metrics, "data_quality_notes", "") or ""
    is_partial = any(kw in notes.upper() for kw in ("PARCIAIS", "PARCIAL", "Q1", "Q2", "Q3", "TRIMESTRE"))
    if is_partial:
        warn_table = doc.add_table(rows=1, cols=1)
        warn_table.style = "Table Grid"
        _set_cell_bg(warn_table.rows[0].cells[0], WARN_BG_HEX)
        warn_p = warn_table.rows[0].cells[0].paragraphs[0]
        warn_p.paragraph_format.space_before = Pt(4)
        warn_p.paragraph_format.space_after  = Pt(4)
        r_icon = warn_p.add_run("⚠  DADOS PARCIAIS — DOCUMENTO PRELIMINAR\n")
        r_icon.bold = True
        r_icon.font.size = Pt(10)
        r_icon.font.color.rgb = RGBColor(0x85, 0x64, 0x04)
        r_body = warn_p.add_run(
            f"Os dados utilizados referem-se a {metrics.data_period}. "
            "Para emissão do PRCE definitivo é necessário o ano completo de dados. "
            "Os valores apresentados são indicativos."
        )
        r_body.font.size = Pt(9)
        r_body.font.color.rgb = RGBColor(0x85, 0x64, 0x04)
        doc.add_paragraph()

    # ── SECTION 1 — INTRODUÇÃO ───────────────────────────────────────────────
    _add_section_heading(doc, "1", "INTRODUÇÃO")
    intro_text = (
        f"O Plano de Racionalização do Consumo de Energia estabelece obrigatoriamente a meta de "
        f"redução dos consumos específicos de energia para a empresa e abrange o período de 3 anos. "
        f"(Portaria 228/90 de 27 de março, artigos 15.º e 16.º). "
        f"A meta não pode ser inferior ao valor calculado pela fórmula:\n\n"
        f"                              M = (C – K) / 2 × n / 3\n\n"
        f"em que M é a meta de redução do consumo específico, C o consumo específico global no "
        f"último ano, e K o valor limite inferior (90% de C)."
    )
    p = doc.add_paragraph(intro_text)
    p.paragraph_format.space_before = Pt(6)

    # ── SECTION 2 — PERÍODO ──────────────────────────────────────────────────
    _add_section_heading(doc, "2", f"PERÍODO DE TRÊS ANOS {plan_start}-{plan_end}")
    doc.add_paragraph(
        f"O período de vigência do presente Plano de Racionalização do Consumo de Energia é de "
        f"três anos, sendo {plan_start} o primeiro ano."
    )

    # ── SECTION 3 — FONTES DE ENERGIA ────────────────────────────────────────
    _add_section_heading(doc, "3", "CARACTERÍSTICAS DAS FONTES DE ENERGIA")
    doc.add_paragraph(
        "Na tabela seguinte são apresentados os valores de conversão de unidades das fontes de "
        "energia e o fator de emissão de CO₂."
    )

    _add_table_caption(doc, "1", "Fatores de conversão")
    t1 = doc.add_table(rows=4, cols=3)
    t1.style = "Table Grid"
    _header_row(t1, ["Fonte de Energia", "Fator de Conversão", "Fator de Emissão CO₂"])
    _data_row(t1, 1, ["Gasóleo", "0,873 kgep/L", "3 098,20 kgCO₂e/tep"], bold_first=True)
    _data_row(t1, 2, ["Gasolina", "0,773 kgep/L", "2 897,30 kgCO₂e/tep"], bold_first=True)
    _data_row(t1, 3, ["Energia Elétrica", "0,290 kgep/kWh", "0,47 kgCO₂e/kWh"], bold_first=True)

    # ── SECTION 4 — CEE NO ANO DE REFERÊNCIA ─────────────────────────────────
    _add_section_heading(doc, "4", f"CONSUMO ESPECÍFICO DE ENERGIA NO ANO DE REFERÊNCIA")
    doc.add_paragraph(
        f"Na tabela seguinte apresentam-se os valores dos consumos específicos verificados "
        f"durante o período {metrics.data_period}."
    )

    _add_table_caption(doc, "2", "Consumo específico de energia – Global")

    has_cargo = metrics.total_cargo_t is not None
    cee_decimals = 5 if metrics.cee_unit == "gep/TK" else 2
    if has_cargo:
        cols = ["", "Distância\nPercorrida (km)", "Consumo\nEnergia (tep)",
                "Carga\nTransportada (t)", "Trabalho\nRealizado (t.km)", f"CEE\n({metrics.cee_unit})"]
    else:
        cols = ["", "Distância\nPercorrida (km)", "Consumo\nEnergia (tep)", f"CEE\n({metrics.cee_unit})"]

    n_rows = len(metrics.quarters) + 2  # header + quarters + total
    t2 = doc.add_table(rows=n_rows, cols=len(cols))
    t2.style = "Table Grid"
    _header_row(t2, cols)

    for i, q in enumerate(metrics.quarters):
        if has_cargo:
            vals = [
                q.label,
                _fmt_num(q.km, 0),
                _fmt_num(q.energy_tep, 1),
                _fmt_num(q.cargo_t, 1),
                _fmt_num(q.work_tkm, 0),
                _fmt_num(q.cee_gep_tk, cee_decimals),
            ]
        else:
            vals = [
                q.label,
                _fmt_num(q.km, 0),
                _fmt_num(q.energy_tep, 1),
                _fmt_num(q.cee_gep_vk, cee_decimals),
            ]
        _data_row(t2, i + 1, vals, bold_first=True, bg_hex=LIGHT_BG_HEX if i % 2 == 0 else None)

    # Total row
    if has_cargo:
        total_vals = [
            "Total",
            _fmt_num(metrics.total_km, 0),
            _fmt_num(metrics.total_energy_tep, 1),
            _fmt_num(metrics.total_cargo_t, 1),
            _fmt_num(metrics.total_work_tkm, 0),
            _fmt_num(metrics.cee_global, cee_decimals),
        ]
    else:
        total_vals = [
            "Total / Média",
            _fmt_num(metrics.total_km, 0),
            _fmt_num(metrics.total_energy_tep, 1),
            _fmt_num(metrics.cee_global, cee_decimals),
        ]
    _data_row(t2, len(metrics.quarters) + 1, total_vals, bold_first=True, bg_hex=TEAL_HEX)
    # Fix total row text color to white
    for cell in t2.rows[len(metrics.quarters) + 1].cells:
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.color.rgb = WHITE

    # ── SECTION 5 — METAS ────────────────────────────────────────────────────
    _add_section_heading(doc, "5", "METAS DE REDUÇÃO DO CONSUMO ESPECÍFICO DE ENERGIA (CEE)")
    doc.add_paragraph("Na tabela seguinte apresentam-se as metas de redução do consumo específico.")

    _add_table_caption(doc, "3", "Metas de redução do consumo específico de energia")
    t3 = doc.add_table(rows=2, cols=6)
    t3.style = "Table Grid"
    _header_row(t3, [
        "Distância\npercorrida (km)",
        "Consumo de\nEnergia (gep)",
        f"CEE\n({metrics.cee_unit})",
        "Valor de K",
        "Redução anual\ndo CEE",
        f"CEE a atingir ao\nfim de 3 anos"
    ])
    _data_row(t3, 1, [
        _fmt_num(metrics.total_km, 0),
        _fmt_num(metrics.total_energy_tep * 1_000_000, 0),
        _fmt_num(metrics.cee_global, cee_decimals),
        _fmt_num(metrics.K, cee_decimals),
        _fmt_num(metrics.annual_reduction, cee_decimals),
        _fmt_num(metrics.cee_global - 3 * metrics.annual_reduction, cee_decimals),
    ])

    # ── SECTION 6 — OBJETIVOS ANUAIS ─────────────────────────────────────────
    _add_section_heading(doc, "6", "OBJETIVOS DOS CONSUMOS ESPECÍFICOS DE ENERGIA ANUAIS")
    doc.add_paragraph("Os objetivos de redução do consumo específico de energia são apresentados na tabela seguinte.")

    _add_table_caption(doc, "4", "Objetivos de redução do consumo específico de energia")
    t4 = doc.add_table(rows=len(metrics.targets) * 2 + 1, cols=2)
    t4.style = "Table Grid"
    _header_row(t4, ["Período", "Frota Global"])
    row_idx = 1
    for year, vals in metrics.targets.items():
        _data_row(t4, row_idx, [f"Ano {year}", f"CEE [{metrics.cee_unit}]"], bold_first=True, bg_hex=LIGHT_BG_HEX)
        _data_row(t4, row_idx + 1, ["", f"{_fmt_num(vals['cee'], cee_decimals)}"])
        row_idx += 2

    # ── SECTION 7 — MEDIDAS ──────────────────────────────────────────────────
    _add_section_heading(doc, "7", "MEDIDAS DE UTILIZAÇÃO RACIONAL DE ENERGIA A SEREM IMPLEMENTADAS")
    doc.add_paragraph(
        "Na tabela seguinte apresenta-se um resumo das medidas de racionalização do consumo "
        f"de energia a implementar no período {plan_start} a {plan_end}."
    )

    _add_table_caption(doc, "5", "Medidas de racionalização do consumo a implementar")
    t5 = doc.add_table(rows=3, cols=8)
    t5.style = "Table Grid"
    _header_row(t5, ["Medida", f"{plan_start}", f"{plan_start+1}", f"{plan_end}",
                     "Total (tep)", "%", "Investimento (€)", "Retorno (anos)"])
    _data_row(t5, 1, ["1. Formação em Eco-Condução",
                      _fmt_num(metrics.total_energy_tep * 0.017, 1),
                      _fmt_num(metrics.total_energy_tep * 0.017, 1),
                      _fmt_num(metrics.total_energy_tep * 0.017, 1),
                      _fmt_num(metrics.total_energy_tep * 0.05, 1),
                      "5,0%", "15 000", "0,2"])
    _data_row(t5, 2, ["Total",
                      _fmt_num(metrics.total_energy_tep * 0.017, 1),
                      _fmt_num(metrics.total_energy_tep * 0.017, 1),
                      _fmt_num(metrics.total_energy_tep * 0.017, 1),
                      _fmt_num(metrics.total_energy_tep * 0.05, 1),
                      "5,0%", "15 000", "0,2"],
              bold_first=True, bg_hex=LIGHT_BG_HEX)

    # ── SECTION 8 — CONSUMOS PREVISTOS ───────────────────────────────────────
    _add_section_heading(doc, "8", "CONSUMOS ESPECÍFICOS PREVISTOS COM A IMPLEMENTAÇÃO DAS MEDIDAS")
    doc.add_paragraph(
        "Na tabela seguinte apresentam-se as reduções previstas para a frota, "
        "tendo em conta as medidas a implementar."
    )

    _add_table_caption(doc, "6", "Previsão da redução do consumo específico de energia")
    t6 = doc.add_table(rows=len(metrics.targets) * 2 + 1, cols=2)
    t6.style = "Table Grid"
    _header_row(t6, ["Período", "Frota Global"])
    row_idx = 1
    for year, vals in metrics.targets.items():
        _data_row(t6, row_idx,     [f"Ano {year}", f"CEE previsto [{metrics.cee_unit}]"],
                  bold_first=True, bg_hex=LIGHT_BG_HEX)
        _data_row(t6, row_idx + 1, ["", f"{_fmt_num(vals['cee'], cee_decimals)}"])
        row_idx += 2

    # ── SECTION 9 — CONCLUSÕES ───────────────────────────────────────────────
    _add_section_heading(doc, "9", "CONCLUSÕES")
    reduction_pct = round(3 * metrics.annual_reduction / metrics.cee_global * 100, 2) if metrics.cee_global > 0 else 5.0
    conclusion = (
        f"Com a implementação do Plano de Racionalização do Consumo de Energia apresentado será "
        f"possível cumprir o disposto no RGCEST — Regulamento da Gestão do Consumo de Energia para "
        f"o Sector dos Transportes, prevendo-se com a implementação das medidas indicadas, uma "
        f"redução de {_fmt_num(reduction_pct, 2)}% do consumo específico de energia.\n\n"
        f"Consumo de energia no ano de referência ({ref_year}): "
        f"{_fmt_num(metrics.total_energy_tep, 1)} tep\n"
        f"Redução prevista: {_fmt_num(metrics.total_energy_tep * reduction_pct / 100, 1)} tep\n"
        f"Consumo previsto no final do período: "
        f"{_fmt_num(metrics.total_energy_tep * (1 - reduction_pct / 100), 1)} tep\n\n"
        f"Atendendo aos factos referidos, solicita-se a aprovação deste Plano de Racionalização, "
        f"para a frota da empresa {company_name}.\n\n\n"
        f"O TÉCNICO RECONHECIDO"
    )
    doc.add_paragraph(conclusion)

    # ── FOOTER NOTE ──────────────────────────────────────────────────────────
    doc.add_paragraph()
    note = doc.add_paragraph()
    run = note.add_run(
        f"⚡ Documento gerado automaticamente por IA em {today} | "
        f"Dados: {metrics.data_period} | "
        f"Motor: Claude Opus 4.8 + PRCE Engine v1.0"
    )
    run.font.size = Pt(7)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if output_path is None:
        safe_name = company_name.replace(" ", "_").replace("/", "-")
        output_path = f"PRCE_{ref_year}_{safe_name}.docx"

    doc.save(output_path)
    return output_path
